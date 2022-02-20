import logging
from io import TextIOWrapper
import os
import docx
from _data import data

def return_breakline(width):
    """Returns a line of ------- corresponding to the width of the logbox in order to visually break text apart"""
    width = width-24  # 20 since padx=10, 4 because... well, I don't know, some weird default padding in the textbox
    count = round(width/8)  # One - character takes up 8 pixels, hence pixels/8
    return '\n' + '-'*count + '\n'


def placeholder():
    print("This is a placeholder")


def get_data_str(data, format_is_selected="{}.*{}", format_not_selected="{}. {}"):
    string = ''
    ALPHABET = 'abcdefghijklmnopqrstuvwxyz'

    for i, id in enumerate(data['question_order']):
        question = data[id]
        index = i + 1  # Since enumerate starts at 0

        # Prints question
        string += "{}. {}\n".format(index, question['question'])
        
        for i, answer in enumerate(question['answers']):
            # Prints answer as two different strings depending on whether answer was selected or not
            if question['is_selected'][i] is True:
                string += "{}.*{}\n".format(ALPHABET[i], answer)
            else:
                string += "{}. {}\n".format(ALPHABET[i], answer)
            
        string += '\n'
    
    return string[:-2]  # Remove last \n


def save_to_doc(file: TextIOWrapper, selected_answers_star=False, selected_answers_bold=True):
    """
    Saves all question to a doc.
    Does so by 
    """
    # Close file since docx needs metadata out of docx files to read them in.
    path = file.name
    file.close()

    doc = docx.Document()
    
    ALPHABET = 'abcdefghijklmnopqrstuvwxyz'
    for i_question, id in enumerate(data['question_order']):
        question = data[id]
        paragraph = doc.add_paragraph()
        
        # Add question and underline it
        run = paragraph.add_run(text=f"{i_question + 1}. {question['question']}\n")
        run.underline = True

        for i_answer, answer in enumerate(question['answers']):
            letter = ALPHABET[i_answer]
            
            if selected_answers_star and question['is_selected'][i_answer]:
                star_char = '*'
            else:
                star_char = ' '
            
            if selected_answers_bold:
                # Creates a run just for the answer letter, then turns it bold
                run = paragraph.add_run(text=letter + '.' + star_char)
                run.bold = True

                # Handles the rest of the answer
                paragraph.add_run(text=answer+'\n')
            else:
                paragraph.add_run(text=f"{letter}.{star_char} {answer}\n")

    doc.save(path)